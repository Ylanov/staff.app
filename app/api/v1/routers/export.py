# app/api/v1/routers/export.py

import io
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, selectinload, joinedload

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.db.database import get_db
from app.models.event import Event, Group, Slot
from app.models.user import User
from app.models.combat_calc import CombatCalcInstance
from app.api.dependencies import get_current_active_admin
from app.api.v1.routers.settings import get_setting

router = APIRouter()


# ─── Вспомогательные функции ──────────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    """Устанавливает границы ячейки. kwargs: top, bottom, left, right — dict с ключами val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        params = kwargs.get(edge)
        if params is None:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), params.get("val", "single"))
        el.set(qn("w:sz"), str(params.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), params.get("color", "000000"))
        tcBorders.append(el)

    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill: str):
    """Заливка ячейки. fill — hex без #, например '4F81BD'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, italic=False,
               size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
               font_name="Times New Roman"):
    """Очищает ячейку и пишет текст с нужным форматированием."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size_pt)


def _merge_row(table, row_idx: int, text: str,
               bold=True, size_pt=10, shading: str = None):
    """Объединяет все ячейки строки и пишет текст — заголовок группы."""
    row = table.rows[row_idx]
    cells = row.cells
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)

    if shading:
        _set_cell_shading(merged, shading)

    _cell_text(merged, text, bold=bold, size_pt=size_pt,
               align=WD_ALIGN_PARAGRAPH.CENTER)


def _thin_border():
    return {"val": "single", "sz": 4, "color": "000000"}


def _apply_row_borders(row):
    for cell in row.cells:
        _set_cell_border(cell,
                         top=_thin_border(), bottom=_thin_border(),
                         left=_thin_border(), right=_thin_border())


def _set_col_width(table, col_idx: int, width_cm: float):
    """Задаёт ширину колонки для всех ячеек."""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ─── Основной роутер (Обычные списки) ─────────────────────────────────────────

@router.get(
    "/events/{event_id}/export-word",
    summary="Выгрузить список в Word (.docx)",
)
def export_event_word(
        event_id: int,
        db: Session = Depends(get_db),
        current_admin: User = Depends(get_current_active_admin),
):
    event = db.query(Event).filter(Event.id == event_id).first()
    if not event:
        raise HTTPException(status_code=404, detail="Событие не найдено")

    # Читаем актуальные данные дежурного из БД — они меняются каждый день
    ORG_NAME = get_setting(db, "org_name")
    DUTY_TITLE = get_setting(db, "duty_title")
    DUTY_RANK = get_setting(db, "duty_rank")
    DUTY_NAME = get_setting(db, "duty_name")

    groups = (
        db.query(Group)
        .filter(Group.event_id == event_id)
        .options(selectinload(Group.slots).joinedload(Slot.position))
        .order_by(Group.order_num)
        .all()
    )

    # ── Собираем данные ───────────────────────────────────────────────────────
    prepared_groups = []
    global_index = 1

    for group in groups:
        slots_data = []
        for slot in sorted(group.slots, key=lambda s: s.id):
            slots_data.append({
                "index": global_index,
                "rank": slot.rank or "",
                "full_name": slot.full_name or "",
                "doc_number": slot.doc_number or "",
                "position": slot.position.name if slot.position else "",
                "callsign": slot.callsign or "",
                "note": slot.note or "",
            })
            global_index += 1
        prepared_groups.append({"name": group.name, "slots": slots_data})

    # ── Создаём документ ──────────────────────────────────────────────────────
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.0)

    # ── Заголовок документа ───────────────────────────────────────────────────
    def _title_para(text: str, bold=False, size_pt=12,
                    space_after_pt=0, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after_pt)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(size_pt)
        return p

    _title_para("Состав", bold=True, size_pt=13)
    _title_para(event.title, bold=True, size_pt=12)
    date_str = event.date.strftime("%d.%m.%Y") if event.date else ""
    _title_para(f"на {date_str}" if date_str else "", bold=False,
                size_pt=11, space_after_pt=4)

    # ── Таблица ───────────────────────────────────────────────────────────────
    COL_WIDTHS = [1.0, 3.2, 4.8, 3.0, 3.2, 2.5, 2.8]
    HEADERS = [
        "№\nп/п",
        "Воинское\nзвание",
        "Фамилия\nИмя Отчество",
        "№\nдокумента",
        "Должность",
        "Позывной",
        "Примечание",
    ]

    total_rows = 1
    for g in prepared_groups:
        total_rows += 1 + len(g["slots"])

    table = doc.add_table(rows=total_rows, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for ci, w in enumerate(COL_WIDTHS):
        _set_col_width(table, ci, w)

    # ── Строка заголовков ─────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for ci, htext in enumerate(HEADERS):
        cell = hdr_row.cells[ci]
        _set_cell_shading(cell, "D9D9D9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_text(cell, htext, bold=True, size_pt=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_border(cell,
                         top=_thin_border(), bottom=_thin_border(),
                         left=_thin_border(), right=_thin_border())

    hdr_row.height = Cm(1.2)

    # ── Строки данных ─────────────────────────────────────────────────────────
    current_row = 1

    for group in prepared_groups:
        g_row = table.rows[current_row]
        g_row.height = Cm(0.65)
        _merge_row(table, current_row, group["name"],
                   bold=True, size_pt=10, shading="E2EFDA")
        _apply_row_borders(g_row)
        current_row += 1

        for slot in group["slots"]:
            row = table.rows[current_row]
            row.height = Cm(0.7)

            values = [
                str(slot["index"]),
                slot["rank"],
                slot["full_name"],
                slot["doc_number"],
                slot["position"],
                slot["callsign"],
                slot["note"],
            ]
            aligns = [
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.LEFT,
            ]

            for ci, (val, align) in enumerate(zip(values, aligns)):
                cell = row.cells[ci]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _cell_text(cell, val, bold=False, size_pt=10, align=align)
                _set_cell_border(cell,
                                 top=_thin_border(), bottom=_thin_border(),
                                 left=_thin_border(), right=_thin_border())

            current_row += 1

    # ── Подпись ───────────────────────────────────────────────────────────────
    doc.add_paragraph()

    if DUTY_TITLE or DUTY_RANK or DUTY_NAME:
        parts = [p for p in [DUTY_TITLE, ORG_NAME, DUTY_RANK] if p]
        left_text = "  ".join(parts)
        right_text = DUTY_NAME

        sign_para = doc.add_paragraph()
        sign_para.paragraph_format.space_before = Pt(4)
        sign_para.paragraph_format.space_after = Pt(0)

        run_left = sign_para.add_run(left_text)
        run_left.font.name = "Times New Roman"
        run_left.font.size = Pt(11)

        run_line = sign_para.add_run("  ___________________  ")
        run_line.font.name = "Times New Roman"
        run_line.font.size = Pt(11)

        run_right = sign_para.add_run(right_text)
        run_right.bold = True
        run_right.font.name = "Times New Roman"
        run_right.font.size = Pt(11)

    # ── Сохраняем в поток ─────────────────────────────────────────────────────
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    safe_title = event.title.replace(" ", "_")
    ascii_name = "Spisok_export.docx"
    unicode_name = quote(f"Список_{safe_title}.docx", safe="")

    headers = {
        "Content-Disposition": (
            f'attachment; filename="{ascii_name}"; '
            f"filename*=UTF-8''{unicode_name}"
        )
    }

    return StreamingResponse(
        file_stream,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers=headers,
    )


# ─── Экспорт Боевого расчёта ──────────────────────────────────────────────────

@router.get(
    "/combat/{instance_id}",
    summary="Выгрузить боевой расчёт в Word (.docx)",
)
def export_combat_calc_word(
        instance_id: int,
        db: Session = Depends(get_db),
        current_admin: User = Depends(get_current_active_admin),
):
    inst = db.query(CombatCalcInstance).filter(CombatCalcInstance.id == instance_id).first()
    if not inst:
        raise HTTPException(status_code=404, detail="Экземпляр расчёта не найден")

    slots_map = {(s.row_key, s.slot_index): s for s in inst.slots}
    structure = inst.template.get_structure()

    # ── Создаём документ ──────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.0)

    def _title_para(text: str, bold=False, size_pt=12, space_after_pt=0, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after_pt)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(size_pt)
        return p

    _title_para(inst.template.title, bold=True, size_pt=14)
    date_str = inst.calc_date.strftime("%d.%m.%Y")
    _title_para(f"на {date_str}", bold=False, size_pt=12, space_after_pt=12)

    # ── Обрабатываем секции шаблона ───────────────────────────────────────────
    for section_data in structure.get("sections", []):
        sec_title = section_data.get("title", "")
        if sec_title:
            _title_para(sec_title, bold=True, size_pt=12, space_after_pt=6, align=WD_ALIGN_PARAGRAPH.LEFT)

        rows_data = section_data.get("rows", [])
        if not rows_data:
            continue

        total_rows = 1
        for r in rows_data:
            total_rows += max(len(r.get("slots", [])), 1)

        table = doc.add_table(rows=total_rows, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        COL_WIDTHS = [5.5, 2.0, 4.0, 2.5, 5.0]
        for ci, w in enumerate(COL_WIDTHS):
            _set_col_width(table, ci, w)

        headers = ["Мероприятие / Состав наряда", "Время", "Кто выделяет", "Место / Подразд.", "В/звание, Фамилия И.О."]
        hdr_row = table.rows[0]
        for ci, h in enumerate(headers):
            cell = hdr_row.cells[ci]
            _set_cell_shading(cell, "D9D9D9")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_text(cell, h, bold=True, size_pt=10)
            _set_cell_border(cell, top=_thin_border(), bottom=_thin_border(), left=_thin_border(), right=_thin_border())
        hdr_row.height = Cm(1.0)

        current_row_idx = 1
        for r in rows_data:
            slots = r.get("slots", [])
            if not slots:
                continue

            start_row_idx = current_row_idx

            for si, slot_def in enumerate(slots):
                row = table.rows[current_row_idx]
                row.height = Cm(0.7)

                slot_data = slots_map.get((r.get("key"), slot_def.get("index", 0)))
                person_str = "—"
                if slot_data and slot_data.full_name:
                    rank = f"{slot_data.rank} " if slot_data.rank else ""
                    person_str = f"{rank}{slot_data.full_name}"

                loc_str = slot_def.get("location", "—")
                if not loc_str.strip():
                    loc_str = "—"

                if si == 0:
                    cell0 = row.cells[0]
                    cell1 = row.cells[1]
                    cell2 = row.cells[2]
                    cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    _cell_text(cell0, r.get("label", ""), size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)
                    _cell_text(cell1, r.get("time", ""), size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
                    _cell_text(cell2, r.get("who_provides", ""), size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)

                cell3 = row.cells[3]
                cell4 = row.cells[4]
                cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell4.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                _cell_text(cell3, loc_str, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)
                _cell_text(cell4, person_str, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)

                for cell in row.cells:
                    _set_cell_border(cell, top=_thin_border(), bottom=_thin_border(), left=_thin_border(),
                                     right=_thin_border())

                current_row_idx += 1

            if len(slots) > 1:
                for col_idx in range(3):
                    start_cell = table.cell(start_row_idx, col_idx)
                    end_cell = table.cell(current_row_idx - 1, col_idx)
                    start_cell.merge(end_cell)

        doc.add_paragraph()

    # ── Подпись ───────────────────────────────────────────────────────────────
    ORG_NAME = get_setting(db, "org_name")
    DUTY_TITLE = get_setting(db, "duty_title")
    DUTY_RANK = get_setting(db, "duty_rank")
    DUTY_NAME = get_setting(db, "duty_name")

    if DUTY_TITLE or DUTY_RANK or DUTY_NAME:
        parts = [p for p in [DUTY_TITLE, ORG_NAME, DUTY_RANK] if p]
        left_text = "  ".join(parts)
        right_text = DUTY_NAME

        sign_para = doc.add_paragraph()
        sign_para.paragraph_format.space_before = Pt(12)

        run_left = sign_para.add_run(left_text)
        run_left.font.name = "Times New Roman"
        run_left.font.size = Pt(11)

        run_line = sign_para.add_run("  ___________________  ")
        run_line.font.name = "Times New Roman"
        run_line.font.size = Pt(11)

        run_right = sign_para.add_run(right_text)
        run_right.bold = True
        run_right.font.name = "Times New Roman"
        run_right.font.size = Pt(11)

    # ── Сохраняем в поток ─────────────────────────────────────────────────────
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    safe_title = inst.template.title.replace(" ", "_").replace('"', '').replace("«", "").replace("»", "")
    ascii_name = "Boevoy_raschet.docx"
    unicode_name = quote(f"Боевой_расчёт_{safe_title}.docx", safe="")

    headers = {
        "Content-Disposition": (
            f'attachment; filename="{ascii_name}"; '
            f"filename*=UTF-8''{unicode_name}"
        )
    }

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )